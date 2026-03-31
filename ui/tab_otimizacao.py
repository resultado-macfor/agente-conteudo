import datetime
import io
import re
import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.content_utils import construir_contexto, realizar_busca_web_perplexity


_LINK_SETA_RE = re.compile(r'(.+?)\s*→\s*(https?://\S+)')


def _normalizar_links(texto: str) -> str:
    def _substituir(m):
        label = m.group(1).strip().lstrip('- ').strip()
        url = m.group(2).strip()
        bullet = '- ' if m.group(0).startswith('- ') else ''
        return f"{bullet}[{label}]({url})"
    return _LINK_SETA_RE.sub(_substituir, texto)


def _adicionar_hyperlink(paragrafo, texto: str, url: str):
    r_id = paragrafo.part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(u)
    run_elem.append(rPr)

    t = OxmlElement('w:t')
    t.text = texto
    run_elem.append(t)
    hyperlink.append(run_elem)
    paragrafo._p.append(hyperlink)


_TOKEN_RE = re.compile(
    r'(\*\*(.+?)\*\*)'
    r'|(\[([^\]]+)\]\((https?://[^\)\s]+)\))'
)


def _adicionar_runs(paragrafo, texto: str):
    pos = 0
    for m in _TOKEN_RE.finditer(texto):
        if m.start() > pos:
            paragrafo.add_run(texto[pos:m.start()])
        if m.group(1):
            paragrafo.add_run(m.group(2)).bold = True
        elif m.group(3):
            _adicionar_hyperlink(paragrafo, m.group(4), m.group(5))
        pos = m.end()
    if pos < len(texto):
        paragrafo.add_run(texto[pos:])


def _adicionar_paragrafo_com_links(doc, texto: str):
    texto = _normalizar_links(texto)
    if not _TOKEN_RE.search(texto):
        doc.add_paragraph(texto)
        return
    p = doc.add_paragraph()
    _adicionar_runs(p, texto)


_HTML_TAG_RE = re.compile(r'<[^>]+>')
_HTML_ENTITY_MAP = {'&amp;': '&', '&lt;': '<', '&gt;': '>', '&nbsp;': ' ', '&quot;': '"', '&apos;': "'"}


def _limpar_html(texto: str) -> str:
    texto = _HTML_TAG_RE.sub('', texto)
    for ent, char in _HTML_ENTITY_MAP.items():
        texto = texto.replace(ent, char)
    return texto


def _gerar_docx(texto: str) -> bytes:
    doc = Document()
    for linha in texto.split("\n"):
        linha = _limpar_html(linha.rstrip())
        if linha.startswith("#### "):
            doc.add_heading(linha[5:], level=4)
        elif linha.startswith("### "):
            doc.add_heading(linha[4:], level=3)
        elif linha.startswith("## "):
            doc.add_heading(linha[3:], level=2)
        elif linha.startswith("# "):
            doc.add_heading(linha[2:], level=1)
        else:
            _adicionar_paragrafo_com_links(doc, linha)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_PROMPT_AVALIACAO = """Você é um editor sênior especializado em SEO para o agronegócio, revisando conteúdos para o portal Mais Agro da Syngenta.

Analise o CONTEÚDO ORIGINAL abaixo e gere um relatório detalhado de melhorias necessárias, avaliando CADA critério da lista abaixo.

{contexto_bloco}
---
{briefing_bloco}

---
{conteudo_bloco}

---
## CRITÉRIOS DE AVALIAÇÃO — verifique cada um:

### 1. METADADOS SEO
- Existe META TITLE (até 60 caracteres com KW principal)?
- Existe META DESCRIPTION (até 155 caracteres com KW e CTA)?
- Existe URL slug amigável, CATEGORIA e ALT TEXT CAPA?

### 2. ESTRUTURA E HEADINGS
- O H1 corresponde exatamente ao título do briefing?
- As seções H2/H3 seguem a estrutura do briefing, na ordem correta?
- Há headings extras não previstos no briefing?

### 3. INTRODUÇÃO
- A introdução começa com dado factual ou afirmação direta?
- Contém frases genéricas proibidas ("No dinâmico cenário...", "Neste artigo vamos...")?
- Apresenta o leitor explicitamente ("Se você é produtor...")?

### 4. QUALIDADE E PROFUNDIDADE
- O conteúdo é raso ou superficial? Falta aprofundamento técnico?
- Os produtos Syngenta são qualificados adequadamente (nome comercial, registro, modo de ação, benefícios)?
- Há informações repetidas entre seções?
- Os parágrafos têm mais de 3 frases ou frases muito longas (>20 palavras)?
- Listas de 3+ itens usam bullet points?

### 5. NEGRITO E FORMATAÇÃO
- A palavra-chave principal está em negrito em MAIS de uma ocorrência? (deve aparecer em negrito APENAS na 1ª vez)
- O negrito é usado em excesso ou em contextos inadequados?
- Existem asteriscos literais `**palavra**` aparecendo como texto (negrito não renderizado)?
- Há tags HTML no texto (`<strong>`, `<br>`, `<p>`, `<a>`, etc.)?

### 6. LINKS INTERNOS (Mais Agro)
- Existem exatamente {qtd_internos} links internos `maisagro.syngenta.com.br`?
- Estão ancorados nos parágrafos onde o tema é discutido (não agrupados no final)?
- Os textos âncora são descritivos e relevantes?

### 7. LINKS EXTERNOS
- Existem exatamente {qtd_externos} links externos de fontes neutras (Embrapa, universidades, institutos)?
- Há links para concorrentes da Syngenta (BASF, Bayer, Corteva, FMC, UPL, Adama, Helm, Nufarm)?
- Há URLs cruas no texto ou formato `texto → URL` com seta?
- Há citações numéricas estilo Wikipedia `[1]`, `[2]`?

### 8. CTA FINAL
- A CTA do briefing está presente e com o texto exato?
- O link da CTA está ancorado corretamente?
- A última frase referencia a central Mais Agro?

---
## FORMATO DO RELATÓRIO:

Para cada critério com problema, escreva:
**[CRITÉRIO]** — Problema identificado: [descrição objetiva do que está errado]
Correção necessária: [o que exatamente deve ser feito]

Ao final, escreva:
**RESUMO:** X problemas críticos, Y melhorias recomendadas.
**PRIORIDADE ALTA:** [lista dos 3 problemas mais impactantes para SEO e qualidade]
"""

_PROMPT_GERACAO = """Você é um especialista em SEO e redação técnica para o agronegócio, escrevendo para o portal Mais Agro da Syngenta.

Reescreva e otimize o CONTEÚDO ORIGINAL aplicando TODOS os ajustes indicados no RELATÓRIO DE AVALIAÇÃO, seguindo o BRIEFING como referência estrutural.

{contexto_bloco}
---
{briefing_bloco}

---
{conteudo_bloco}

---
###RELATÓRIO DE AVALIAÇÃO (aplique TODOS os pontos)###
{avaliacao}
###FIM DO RELATÓRIO###

{fontes_bloco}

---
## INSTRUÇÕES DE GERAÇÃO

### 1. METADADOS SEO — coloque SEMPRE no topo, antes do H1
```
META TITLE: [até 60 caracteres, com KW principal]
META DESCRIPTION: [até 155 caracteres, com KW e chamada para ação]
URL: /[slug-amigavel-baseado-no-h1]
CATEGORIA: [categoria sugerida]
ALT TEXT CAPA: [texto descritivo com KW]
```

### 2. ESTRUTURA DO ARTIGO
- H1 exatamente como indicado no briefing
- Heading de partida do corpo: **{nivel_heading}**
- Seções na ordem exata do briefing — nada a mais, nada a menos
- Cada seção: 2 a 4 parágrafos + bullets quando há 3+ itens paralelos
- Bullets: `- **Termo:** explicação da característica`
- Listas numeradas para sequências de etapas (1. 2. 3.)

### 3. INTRODUÇÃO
- Comece com dado factual ou afirmação direta sobre o tema
- PROIBIDO: aberturas genéricas, apresentar o leitor, frases meta

### 4. QUALIDADE TEXTUAL
- Parágrafos com NO MÁXIMO 3 frases curtas (~20 palavras)
- KW principal em negrito SOMENTE na 1ª ocorrência — nas demais, sem negrito
- Cada informação aparece UMA única vez — sem repetição entre seções
- Qualifique os produtos Syngenta: nome comercial + registro + modo de ação + benefícios

### 5. LINKS
**Formato obrigatório:** `[texto descritivo](https://url-completa.com)`

**PROIBIDO:** URLs cruas, formato `texto → URL`, citações `[n]`, concorrentes Syngenta (BASF, Bayer, Corteva, FMC, UPL, Adama, Helm, Nufarm)

**Internos (Mais Agro) — exatamente {qtd_internos} distribuídos no corpo:**
- Ancorados no parágrafo onde o tema é discutido
- Padrão: `https://maisagro.syngenta.com.br/[slug]`
- Exemplo: O [manejo integrado de plantas daninhas](https://maisagro.syngenta.com.br/manejo-plantas-daninhas) reduz o banco de sementes...

**Externos — exatamente {qtd_externos} distribuídos no corpo:**
- Apenas Embrapa, universidades, institutos governamentais
- Ancorados no dado/estudo: `De acordo com [pesquisa da Embrapa](https://embrapa.br/...), a espécie...`

### 6. PRODUTOS SYNGENTA
- Seção própria antes da CTA com heading específico
- Mencione nome comercial + registro (ex: CALARIS®, Dual Gold®, Grover®)
- Tom: "A Syngenta oferece qualidade e tecnologia..."

### 7. CTA FINAL
- Texto exato do briefing
- Última frase: "Confira a central de conteúdos [Mais Agro](URL-da-CTA) para ficar por dentro de tudo o que está acontecendo no campo."

### 8. FORMATAÇÃO — PROIBIÇÕES ABSOLUTAS
- NUNCA tags HTML: `<strong>`, `<em>`, `<b>`, `<i>`, `<a>`, `<br>`, `<p>`, `<ul>`, `<li>`
- NUNCA asteriscos literais como texto — use Markdown: `**negrito**`
- Markdown puro em todo o artigo

---
**LEMBRETE:** Markdown puro, zero HTML, zero `[n]`, zero concorrentes, links ancorados nos parágrafos, CTA com link no final.

Gere o artigo completo.
"""


def render(tab, modelo_texto):
    with tab:
        st.header("🚀 Otimização SEO de Conteúdo")

        for key, default in [
            ('otimizacao_avaliacao', None),
            ('otimizacao_conteudo_final', None),
            ('otimizacao_briefing', ""),
            ('otimizacao_conteudo_original', ""),
            ('otimizacao_fontes', ""),
            ('ajustes_realizados', []),
        ]:
            if key not in st.session_state:
                st.session_state[key] = default

        modo_entrada = st.radio(
            "Modo de entrada:",
            ["Briefing + Conteúdo original", "Apenas Briefing", "Apenas Conteúdo original"],
            horizontal=True,
            help="Escolha quais insumos serão usados na avaliação e geração.",
        )

        col_esq, col_dir = st.columns(2)
        with col_esq:
            usar_briefing = modo_entrada in ("Briefing + Conteúdo original", "Apenas Briefing")
            briefing_entrada = st.text_area(
                "📋 Briefing de entrada:" + ("" if usar_briefing else " (não utilizado)"),
                height=280,
                placeholder="Cole aqui o briefing com título H1, KWs, estrutura H2/H3, CTA, tom, etc.",
                value=st.session_state.otimizacao_briefing,
                disabled=not usar_briefing,
            )
        with col_dir:
            usar_conteudo = modo_entrada in ("Briefing + Conteúdo original", "Apenas Conteúdo original")
            conteudo_original = st.text_area(
                "📄 Conteúdo original:" + ("" if usar_conteudo else " (não utilizado)"),
                height=280,
                placeholder="Cole aqui o conteúdo existente que será avaliado e otimizado.",
                value=st.session_state.otimizacao_conteudo_original,
                disabled=not usar_conteudo,
            )

        col1, col2, col3, col4 = st.columns([2, 1, 1, 1])
        with col1:
            usar_busca_web = st.checkbox("🔍 Enriquecer com busca web (Perplexity)", value=False)
        with col2:
            nivel_heading = st.selectbox(
                "Heading de partida do corpo:",
                ["H2", "H3", "H1"],
                help="Nível do primeiro heading do corpo do artigo (conforme briefing)",
            )
        with col3:
            qtd_links_internos = st.number_input(
                "Links internos",
                min_value=1, max_value=10, value=3, step=1,
                help="Quantidade de links internos Mais Agro a incluir no artigo",
            )
        with col4:
            qtd_links_externos = st.number_input(
                "Links externos",
                min_value=0, max_value=10, value=2, step=1,
                help="Quantidade de links externos (Embrapa, universidades, etc.) a incluir",
            )

    
        st.divider()
        st.subheader("Etapa 1 — Avaliação do Conteúdo")

        if st.button("🔍 Avaliar Conteúdo", type="secondary", use_container_width=True):
            if usar_briefing and not briefing_entrada.strip():
                st.warning("Cole o briefing de entrada.")
                st.stop()
            if usar_conteudo and not conteudo_original.strip():
                st.warning("Cole o conteúdo original.")
                st.stop()
            if not usar_briefing and not usar_conteudo:
                st.warning("Selecione ao menos um modo de entrada.")
                st.stop()

            st.session_state.otimizacao_briefing = briefing_entrada if usar_briefing else ""
            st.session_state.otimizacao_conteudo_original = conteudo_original if usar_conteudo else ""
            st.session_state.otimizacao_avaliacao = None
            st.session_state.otimizacao_conteudo_final = None

            with st.spinner("Analisando conteúdo..."):
                try:
                    contexto_agente = ""
                    if st.session_state.get("agente_selecionado"):
                        agente = st.session_state.agente_selecionado
                        contexto_agente = construir_contexto(agente, st.session_state.get("segmentos_selecionados", []))

                    contexto_bloco = f"###CONTEXTO DO AGENTE###\n{contexto_agente}\n###FIM DO CONTEXTO###\n\n" if contexto_agente else ""

                    briefing_bloco = f"###BRIEFING DE REFERÊNCIA###\n{briefing_entrada}\n###FIM DO BRIEFING###\n" if usar_briefing and briefing_entrada.strip() else "(Briefing não fornecido — avalie apenas o conteúdo.)"
                    conteudo_bloco = f"###CONTEÚDO ORIGINAL###\n{conteudo_original}\n###FIM DO CONTEÚDO ORIGINAL###\n" if usar_conteudo and conteudo_original.strip() else "(Conteúdo original não fornecido — gere a partir do briefing.)"

                    prompt_aval = _PROMPT_AVALIACAO.format(
                        contexto_bloco=contexto_bloco,
                        briefing_bloco=briefing_bloco,
                        conteudo_bloco=conteudo_bloco,
                        qtd_internos=qtd_links_internos,
                        qtd_externos=qtd_links_externos,
                    )

                    resposta = modelo_texto.generate_content(prompt_aval)
                    st.session_state.otimizacao_avaliacao = resposta.text
                    st.success("✅ Avaliação concluída!")

                except Exception as e:
                    st.error(f"❌ Erro na avaliação: {e}")

        if st.session_state.otimizacao_avaliacao:
            st.success("✅ Conteúdo analisado com sucesso!")

            st.divider()
            st.subheader("Etapa 2 — Geração do Conteúdo Otimizado")
            avaliacao_editavel = st.session_state.otimizacao_avaliacao

            if st.button("🚀 Gerar Conteúdo Otimizado", type="primary", use_container_width=True):
                with st.spinner("Gerando conteúdo otimizado..."):
                    try:
                        fontes_encontradas = ""
                        if usar_busca_web:
                            busca_ph = st.empty()
                            try:
                                resultado_busca = realizar_busca_web_perplexity(
                                    st.session_state.otimizacao_conteudo_original, "SEO", "técnico"
                                )
                                if resultado_busca and not resultado_busca.startswith("❌"):
                                    fontes_encontradas = resultado_busca
                                    st.session_state.otimizacao_fontes = resultado_busca
                                    busca_ph.success(f"✅ Busca concluída ({len(resultado_busca.split())} palavras)")
                                else:
                                    busca_ph.warning("⚠️ Busca não retornou resultados válidos")
                            except Exception as e:
                                st.warning(f"⚠️ Erro na busca web: {e}")

                        contexto_agente = ""
                        if st.session_state.get("agente_selecionado"):
                            agente = st.session_state.agente_selecionado
                            contexto_agente = construir_contexto(agente, st.session_state.get("segmentos_selecionados", []))

                        contexto_bloco = f"###CONTEXTO DO AGENTE###\n{contexto_agente}\n###FIM DO CONTEXTO###\n\n" if contexto_agente else ""
                        fontes_bloco = f"---\n###FONTES WEB###\n{fontes_encontradas}\n###FIM DAS FONTES###\n" if fontes_encontradas else ""

                        _briefing_salvo = st.session_state.otimizacao_briefing
                        _conteudo_salvo = st.session_state.otimizacao_conteudo_original
                        briefing_bloco_ger = f"###BRIEFING###\n{_briefing_salvo}\n###FIM DO BRIEFING###\n" if _briefing_salvo else "(Briefing não fornecido — baseie-se no conteúdo original e no relatório.)"
                        conteudo_bloco_ger = f"###CONTEÚDO ORIGINAL###\n{_conteudo_salvo}\n###FIM DO CONTEÚDO ORIGINAL###\n" if _conteudo_salvo else "(Conteúdo original não fornecido — gere a partir do briefing e do relatório.)"

                        prompt_ger = _PROMPT_GERACAO.format(
                            contexto_bloco=contexto_bloco,
                            briefing_bloco=briefing_bloco_ger,
                            conteudo_bloco=conteudo_bloco_ger,
                            avaliacao=avaliacao_editavel,
                            fontes_bloco=fontes_bloco,
                            nivel_heading=nivel_heading,
                            qtd_internos=qtd_links_internos,
                            qtd_externos=qtd_links_externos,
                        )

                        resposta = modelo_texto.generate_content(prompt_ger)
                        st.session_state.otimizacao_conteudo_final = resposta.text
                        st.success("✅ Conteúdo gerado com sucesso!")

                    except Exception as e:
                        st.error(f"❌ Erro na geração: {e}")

        if st.session_state.otimizacao_conteudo_final:
            resultado = st.session_state.otimizacao_conteudo_final
            st.divider()

            if "META TITLE:" in resultado:
                linhas = resultado.split("\n")
                meta_linhas = [l for l in linhas if any(
                    l.startswith(k) for k in ("META TITLE:", "META DESCRIPTION:", "URL:", "CATEGORIA:", "ALT TEXT")
                )]
                if meta_linhas:
                    st.subheader("📊 Metadados SEO")
                    for l in meta_linhas:
                        st.markdown(f"**{l}**")

            st.subheader("📝 Conteúdo Otimizado")
            st.markdown(resultado)

            col_m1, col_m2, col_m3 = st.columns(3)
            with col_m1:
                st.metric("Palavras", len(resultado.split()))
            with col_m2:
                headings = resultado.count("## ") + resultado.count("### ")
                st.metric("Headings", headings)
            with col_m3:
                tem_cta = "maisagro" in resultado.lower() or "syngenta" in resultado.lower()
                st.metric("CTA", "✅" if tem_cta else "❌")

            st.download_button(
                "💾 Baixar Conteúdo Otimizado (.docx)",
                data=_gerar_docx(resultado),
                file_name=f"conteudo_otimizado_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

            st.divider()
            st.subheader("🔄 Ajustes Incrementais")

            comando_ajuste = st.text_area(
                "Descreva os ajustes:",
                height=80,
                placeholder="Ex: Reescreva a introdução, aprofunde a seção X, ajuste o tom...",
                key="ajuste_text",
            )

            if st.button("🔄 Aplicar Ajustes", key="btn_ajuste"):
                if not comando_ajuste.strip():
                    st.warning("Descreva os ajustes desejados.")
                    st.stop()

                with st.spinner("Aplicando ajustes..."):
                    try:
                        prompt_ajuste = f"""Você é um especialista em SEO e redação técnica para o agronegócio, escrevendo para o portal Mais Agro da Syngenta.

Aplique os ajustes solicitados ao conteúdo abaixo, mantendo INTEGRALMENTE:
- Metadados SEO (META TITLE, META DESCRIPTION, URL, CATEGORIA, ALT TEXT) no topo
- Hierarquia de headings existente
- Links internos Mais Agro ancorados nos parágrafos onde o tema é mencionado
- Apenas as seções previstas — sem adicionar seções extras
- Parágrafos com no máximo 3 frases curtas
- Bullets com `**Termo:** explicação`
- Seção de produtos Syngenta recomendados
- CTA final com link âncora para o Mais Agro
- Markdown puro — zero HTML, zero `[n]`, zero concorrentes Syngenta
- KW principal em negrito SOMENTE na 1ª ocorrência

**CONTEÚDO ATUAL:**
{st.session_state.otimizacao_conteudo_final}

**AJUSTES SOLICITADOS:**
{comando_ajuste}

Retorne o conteúdo completo com os ajustes aplicados.
"""
                        resposta = modelo_texto.generate_content(prompt_ajuste)
                        st.session_state.otimizacao_conteudo_final = resposta.text
                        st.session_state.ajustes_realizados.append(comando_ajuste)

                        st.success("✅ Ajustes aplicados!")
                        st.rerun()

                    except Exception as e:
                        st.error(f"Erro: {e}")

            if st.session_state.ajustes_realizados:
                with st.expander(f"Histórico de ajustes ({len(st.session_state.ajustes_realizados)})"):
                    for i, aj in enumerate(st.session_state.ajustes_realizados, 1):
                        st.markdown(f"{i}. {aj}")
                if st.button("🗑️ Limpar histórico"):
                    st.session_state.ajustes_realizados = []
                    st.rerun()
